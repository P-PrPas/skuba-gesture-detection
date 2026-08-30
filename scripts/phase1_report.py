"""Build the Phase 1 backbone test report (.docx) from the eval JSON.

    python scripts/phase1_eval.py --device cpu --annotate
    python scripts/phase1_eval.py --device gpu --annotate
    python scripts/phase1_report.py            # -> results/phase1/backbone_report.docx

Reads results/phase1/metrics_{cpu,gpu}.json and the montages, writes one .docx
with the verdict up top, full latency/VRAM/accuracy tables, and the annotated
montages embedded so the reader can judge keypoint quality directly.
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "results" / "phase1"

POSE_LABEL = {
    "mediapipe_pose": "MediaPipe Pose (Solutions API)",
    "mediapipe_tasks_lite": "MediaPipe Tasks — lite",
    "mediapipe_tasks_full": "MediaPipe Tasks — full",
    "mediapipe_tasks_heavy": "MediaPipe Tasks — heavy",
    "yolo11n-pose": "YOLO11n-pose",
    "yolo11s-pose": "YOLO11s-pose",
    "rtmpose_lightweight": "RTMPose-t (rtmlib)",
    "rtmpose_balanced": "RTMPose-m (rtmlib)",
}

# qualitative keypoint-quality notes from the montage review (section 5)
QUALITY = {
    "mediapipe_pose": "Stable on squat/sit. No flyaway keypoints; single-person model so it never jumps to the bystander. Occluded legs inferred, occasionally implausible on the worst floor frames.",
    "mediapipe_tasks_lite": "Same skeleton quality as the Solutions API — clean on squat/sit, no flyaways, no bystander — but ~1.5x faster on CPU. This is the one to use.",
    "mediapipe_tasks_full": "Marginally steadier than lite on the occluded floor frames; ~1.5x slower.",
    "mediapipe_tasks_heavy": "Best MediaPipe skeleton, but 108 ms/frame on CPU — no faster than YOLO11n and much slower than lite.",
    "yolo11n-pose": "Head/face keypoints repeatedly shoot to the frame corners at high confidence; skeleton jumps to the background bystander in several frames.",
    "yolo11s-pose": "Torso cleaner than 11n but the corner-flyaway head keypoints and bystander pickup remain.",
    "rtmpose_lightweight": "Skeleton usually clean; the bundled YOLOX-tiny detector picks the bystander in a minority of frames.",
    "rtmpose_balanced": "Best skeleton of the lot — clean through the whole deep crouch, no flyaways. Costs: 356 ms/frame on CPU, and the YOLOX-m detector can still grab a bystander.",
}
RECOMMEND = {"mediapipe_tasks_lite": "YES - default"}


def scorecard_rows(cpu, gpu):
    rows = []
    keys = list(cpu["pose"])
    keys = [k for k in keys if RECOMMEND.get(k)] + [k for k in keys if not RECOMMEND.get(k)]
    for key in keys:
        v = cpu["pose"][key]
        g = (gpu or {}).get("pose", {}).get(key, {})
        gv = (g.get("vram") or {}).get("model_mb")
        cpu_ms = v["clips"].get("squat", {}).get("ms_mean", 0)
        gd = g.get("clips", {}).get("squat", {})
        gpu_ms = gd.get("ms_mean")
        gpu_ran = g.get("device_actual", "")
        det = [c.get("detect_pct", 0) for c in v["clips"].values()]
        rows.append([
            POSE_LABEL.get(key, key),
            "0 (CPU)" if not gv else f"{gv:.0f}",
            f"{cpu_ms:.0f}",
            "-" if gpu_ms is None else (f"{gpu_ms:.0f}" + ("" if gpu_ran == "cuda" else " (CPU)")),
            f"{sum(det) / len(det):.0f}%" if det else "-",
            QUALITY.get(key, ""),
            RECOMMEND.get(key, "no"),
        ])
    return rows


def _load(dev):
    p = OUT / f"metrics_{dev}.json"
    return json.loads(p.read_text()) if p.exists() else None


def _cell(table, r, c, text, bold=False):
    cell = table.cell(r, c)
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for c, h in enumerate(headers):
        _cell(t, 0, c, h, bold=True)
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            _cell(t, r, c, val)
    return t


def latency_rows(data):
    rows = []
    for key, v in data["pose"].items():
        sq = v["clips"].get("squat", {})
        la = v["clips"].get("laying", {})
        si = v["clips"].get("sit", {})
        vram = (v.get("vram") or {}).get("model_mb")
        rows.append([
            POSE_LABEL.get(key, key), v["device_actual"],
            f"{sq.get('ms_mean', 0):.1f}", f"{sq.get('ms_p90', 0):.1f}",
            f"{sq.get('fps', 0):.1f}",
            f"{la.get('ms_mean', 0):.1f}", f"{si.get('ms_mean', 0):.1f}",
            "0" if vram == 0 else ("-" if vram is None else f"{vram:.0f}"),
        ])
    return rows


def detect_rows(data):
    clips = ["laying", "squat", "sit"]
    rows = []
    for key, v in data["pose"].items():
        rows.append([POSE_LABEL.get(key, key)] +
                    [v["clips"].get(c, {}).get("detect_rate", "-") for c in clips])
    return rows


def main():
    cpu, gpu = _load("cpu"), _load("gpu")
    if cpu is None:
        raise SystemExit("run scripts/phase1_eval.py --device cpu first")

    doc = Document()
    doc.add_heading("SKUBA gesture detection — Phase 1 backbone test", 0)
    doc.add_paragraph(f"Generated {date.today().isoformat()}").runs[0].italic = True

    h = cpu["host"]
    verdict = doc.add_paragraph()
    r = verdict.add_run("Verdict: MediaPipe Pose (Tasks API, lite model) + MediaPipe Hands.")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1a, 0x7f, 0x37)
    doc.add_paragraph(
        "Switching from the legacy MediaPipe Solutions API to the newer Tasks API "
        "with the 'lite' pose model runs at 25 ms/frame (40 FPS) on CPU — "
        "essentially matching YOLO11n-pose on the GPU (24 ms), with zero VRAM, the "
        "same 33-landmark output, a built-in single-person model (no bystander "
        "lock-on), and none of YOLO's flyaway-keypoint bug. So the GPU question is "
        "largely moot for pose: the fast MediaPipe path is already as fast as the "
        "GPU alternatives and keeps the CUDA devices free for the other modules."
    )
    doc.add_paragraph(
        "On MediaPipe + GPU specifically (what you asked for): the delegate exists "
        "and works on Linux — the deploy OS — but not on Windows, so it cannot be "
        "measured on this dev machine (see section 5c for what to expect and how to "
        "measure it on the Acer). It uses OpenGL ES compute, not CUDA, so it would "
        "not contend with the CUDA modules for compute cores; its VRAM footprint is "
        "small (tens of MB) but non-zero."
    )

    doc.add_heading("Scorecard", 1)
    _table(doc,
           ["Backbone", "VRAM MB", "CPU ms/f", "GPU ms/f", "Detect (avg)", "Keypoint quality", "Use it?"],
           scorecard_rows(cpu, gpu))
    doc.add_paragraph(
        "CPU/GPU ms are per frame on squat_01, this machine (RTX 3050, proxy for "
        "the Acer). VRAM is the CUDA figure for the GPU-capable backends; MediaPipe "
        "is 0 on CPU and small-but-unmeasured on the Linux GL delegate. 'Detect' = "
        "frames with a person found, averaged over the three posture clips."
    ).runs[0].font.size = Pt(9)

    doc.add_heading("1. Test setup", 1)
    doc.add_paragraph(
        f"Host: {h['platform']}, Python {h['python']}. "
        f"GPU: {h.get('gpu') or 'n/a'} ({h.get('gpu_mem_total_mb') or '?'} MB). "
        "Latency is measured on this machine and is only a proxy for the Acer/Ubuntu "
        "deploy laptop; treat the ranking and the VRAM figures as the durable result."
    )
    doc.add_paragraph(
        "Hard-case clips (from data/clips/, subject at ~2–4 m): posture — laying_01, "
        "squat_01, sit_02; hands — rock_01, ok_01, i_love_you_01, two_finger_01. "
        "For each backbone we record per-frame latency (mean / p90 / FPS), the "
        "detection rate, peak VRAM, and an 8-frame annotated montage for visual "
        "keypoint-quality review."
    )
    doc.add_paragraph(
        "MediaPipe is tested two ways: the legacy Solutions API (model_complexity=1, "
        "what backbone/pose.py currently uses) and the current Tasks API "
        "(pose_landmarker lite / full / heavy). Only the Tasks API has a GPU "
        "delegate, and only on Linux."
    )

    doc.add_heading("2. Body-pose latency & VRAM — CPU", 1)
    _table(doc,
           ["Backbone", "Ran on", "squat ms/f", "p90 ms", "FPS", "laying ms/f", "sit ms/f", "VRAM MB"],
           latency_rows(cpu))

    if gpu:
        doc.add_heading("3. Body-pose latency & VRAM — GPU requested", 1)
        _table(doc,
               ["Backbone", "Ran on", "squat ms/f", "p90 ms", "FPS", "laying ms/f", "sit ms/f", "VRAM MB"],
               latency_rows(gpu))
    else:
        doc.add_heading("3. GPU pass", 1)
        doc.add_paragraph("Not available / not run on this machine.")

    doc.add_heading("4. Detection rate (person found / total frames)", 1)
    doc.add_paragraph("CPU pass:")
    _table(doc, ["Backbone", "laying", "squat", "sit"], detect_rows(cpu))

    doc.add_heading("5. Keypoint quality — visual review", 1)
    doc.add_paragraph(
        "8 frames per clip, evenly spaced. Green = skeleton, orange = keypoints. "
        "Look for: keypoints flying to the frame edge, the skeleton jumping to a "
        "second person, limbs snapping to implausible angles. Keypoints are the "
        "same on CPU and GPU (identical models) — only latency changes. Full "
        "annotated clips: results/phase1/annotated/."
    )
    # lead with the recommended one, then the rest
    order = ["mediapipe_tasks_lite"] + [k for k in cpu["pose"] if k != "mediapipe_tasks_lite"]
    for key in order:
        for clip in ("squat", "laying", "sit"):
            img = OUT / "montages" / f"{key}__{clip}__cpu.jpg"
            if img.exists():
                doc.add_paragraph(f"{POSE_LABEL.get(key, key)} — {clip}", style="Heading 3")
                doc.add_picture(str(img), width=Inches(6.5))

    doc.add_heading("5b. The GPU options for pose, laid out", 1)
    doc.add_paragraph(
        "With CUDA already installed on the robot and the other modules already "
        "using the GPU, VRAM is no longer a hard blocker — so here is the honest "
        "comparison of every GPU-capable option against the fast CPU path:"
    )
    _table(doc, ["Option", "Latency (squat)", "VRAM", "Runtime cost", "Keypoint issues"], [
        ["MediaPipe Tasks lite — CPU", "25 ms / 40 FPS", "0", "none (already a dep)", "none"],
        ["MediaPipe Tasks lite — GPU (Linux)", "not measurable on Windows; expect similar-to-somewhat-faster, see 5c", "tens of MB (GL, not CUDA)", "needs an EGL/GL context on the headless robot", "none"],
        ["YOLO11n-pose — CUDA", "24 ms / 42 FPS", "70 MB", "torch + CUDA (~5 GB on disk)", "flyaway head keypoints; picks the bystander — needs a person-selection fix"],
        ["YOLO11s-pose — CUDA", "21 ms / 47 FPS", "130 MB", "torch + CUDA", "same as 11n, milder"],
        ["RTMPose-t — CUDA", "37 ms / 27 FPS", "359 MB", "onnxruntime-gpu + detector", "detector picks the bystander sometimes"],
        ["RTMPose-m — CUDA", "83 ms / 12 FPS", "611 MB", "onnxruntime-gpu + detector", "best skeleton, but slow and bystander-prone"],
    ])
    doc.add_paragraph(
        "The takeaway: the fastest GPU option (YOLO11n at 24 ms) is no faster than "
        "MediaPipe Tasks lite on the CPU (25 ms), and it brings back the two "
        "keypoint bugs plus a person-selection rewrite. There is no pose backbone "
        "here where spending GPU buys a meaningful win. Keep pose on the CPU; leave "
        "the GPU for the modules that genuinely need it."
    )

    doc.add_heading("5c. MediaPipe + GPU — what it is and how to measure it", 1)
    doc.add_paragraph(
        "The MediaPipe Tasks GPU delegate raises NotImplementedError on Windows "
        "(recorded in metrics_gpu.json), so it is not in the tables above. On Linux "
        "— the deploy OS — it is available. Key facts:"
    )
    for t in [
        "It runs on the GPU via OpenGL ES 3.1 compute shaders, NOT CUDA. It does "
        "not use the CUDA runtime and does not allocate through the CUDA allocator, "
        "so it will not compete with the CUDA modules for SM/compute scheduling the "
        "way a second torch model would.",
        "VRAM: the models are 3–30 MB of weights; with GL working buffers the "
        "resident footprint is on the order of tens of MB. It shows up in "
        "nvidia-smi's process list but not in torch.cuda memory stats.",
        "Speed: on desktop-class hardware the XNNPACK CPU path for the lite/full "
        "models is already well optimised; the GPU delegate typically ranges from "
        "roughly on-par to ~2x faster, and can be slower for the lite model because "
        "the per-frame CPU<->GPU image upload dominates. It helps most for the "
        "heavy model.",
        "It needs a usable GL context. On a headless robot that means EGL with a "
        "surfaceless context; if the other modules already hold the GL/EGL display "
        "this is usually fine, otherwise it needs configuring.",
    ]:
        doc.add_paragraph(t, style="List Bullet")
    doc.add_paragraph(
        "To get real numbers on the Acer: install mediapipe on the Ubuntu box and "
        "run  python scripts/phase1_eval.py --device gpu --only mediapipe_tasks "
        "--merge  there. The MPTasksPose backend already requests the GPU delegate "
        "and records whether it loaded; on Linux it will produce a real "
        "'gpu(gl)' row with nvidia-smi VRAM delta."
    )

    doc.add_heading("6. Hands — MediaPipe Hands on wrist-anchored crops", 1)
    hrows = [[c, v["detect_rate"], f"{v['ms_mean']} ms" if v["ms_mean"] else "-"]
             for c, v in cpu["hands"]["clips"].items()]
    _table(doc, ["Gesture clip", "detected / crops", "ms/crop"], hrows)
    doc.add_paragraph(
        "Misses are concentrated at the moments the hand enters/leaves frame "
        "(motion blur), not on the held gesture — which is exactly what the "
        "presence-flag + temporal-smoothing design absorbs."
    )

    doc.add_heading("7. Combined pipeline (the real workload)", 1)
    c = cpu["combined_mediapipe"]
    doc.add_paragraph(
        f"MediaPipe Pose (Solutions API, current backbone/pose.py) + 2 wrist-crop "
        f"hand detections per frame: {c['ms_mean']:.0f} ms/frame ≈ {c['fps']:.1f} FPS "
        "on this CPU. The two hand crops are ~2/3 of that. Swapping pose to the "
        "Tasks lite model saves ~13 ms/frame (pose 38 → 25 ms), lifting the "
        "combined rate to roughly 10 FPS; the hands remain the bottleneck and are "
        "where any further speed work (Phase 5) should go — smaller crops, or "
        "VIDEO-mode tracking instead of per-frame palm detection."
    )

    doc.add_heading("8. Recommendation", 1)
    doc.add_paragraph(
        "Adopt MediaPipe Pose via the Tasks API (pose_landmarker lite) + MediaPipe "
        "Hands. Update backbone/pose.py from the Solutions API to the Tasks API.",
        style="List Bullet")
    doc.add_paragraph(
        "Rationale: (1) 25 ms/frame on CPU — as fast as YOLO11n on the GPU, ~1.5x "
        "faster than the Solutions API we started with; (2) 0 VRAM, keeps the GPU "
        "free for the modules that need CUDA; (3) stable keypoints, single-person "
        "model, no bystander lock-on, no flyaway bug; (4) 33 landmarks (vs COCO-17), "
        "matching features/schema.py; (5) hand landmarks stable on the "
        "overlapping-finger shapes (rock, ILY).", style="List Bullet")
    doc.add_paragraph(
        "GPU was evaluated at your request (sections 5b–5c). Result: no pose "
        "backbone here gets a meaningful speed win from the GPU — the fastest GPU "
        "option (YOLO11n, 24 ms) ties the CPU MediaPipe path and reintroduces two "
        "keypoint bugs. MediaPipe's own GPU delegate is Linux-only (can't measure "
        "on this box) and, being OpenGL-based, would add tens of MB of VRAM for at "
        "best a small gain. Keep pose on CPU.", style="List Bullet")
    doc.add_heading("Open items before Phase 1 is fully signed off", 2)
    doc.add_paragraph("Record a clean 'laying' clip from the robot's camera height and "
                      "re-run — the current laying_01 is mostly transition and occlusion.",
                      style="List Bullet")
    doc.add_paragraph("On the Acer/Ubuntu: run the combined pipeline and confirm the "
                      "real-time rate (~10 FPS expected with Tasks-lite pose); also run "
                      "the MediaPipe GPU delegate there (section 5c) to close the GPU "
                      "question with a real number.", style="List Bullet")
    doc.add_paragraph("Port backbone/pose.py to the Tasks API and re-extract features "
                      "(the 33-landmark order is the same, but confirm — features are "
                      "only valid for the exact backbone, see ARCHITECTURE.md).",
                      style="List Bullet")

    out = OUT / "backbone_report.docx"
    doc.save(str(out))
    print(f"-> {out}")


if __name__ == "__main__":
    main()
