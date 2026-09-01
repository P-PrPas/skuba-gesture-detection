"""Build results/phase2/dataset_report.docx — what external data we pulled, with
sample images, and how the training dataset was constructed.

    python scripts/pull_samples.py            # sample images -> data/samples_ext/
    python -m pipeline.build_dataset          # -> data/dataset/ + card
    python scripts/phase2_report.py
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import cv2
import numpy as np
from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DS = ROOT / "data" / "dataset"
SAMP = ROOT / "results" / "phase2"
OUT = SAMP
SAMPLES = ROOT / "data" / "samples_ext"
GRIDS = SAMP / "grids"

SOURCES = [
    ("HaGRID v1", "cj-mills/hagrid-sample-30k-384p (HF mirror of hukenovs/hagrid)",
     "webcam / phone images, upper body, hand bbox <=16% of frame",
     "custom BY-SA-style, no NC clause (RoboCup use fine)",
     "ok, two_finger (peace), rock, thumb (like); call -> idle hard-negative"),
    ("HaGRID no_gesture", "cj-mills/hagrid-classification-512p-no-gesture-150k",
     "same, hand visible but no gesture", "as above", "idle"),
    ("HaGRID v2", "testdummyvt/hagRIDv2_512px (val split)",
     "same, 34 classes", "disputed commercial status; RoboCup use fine",
     "mini_heart (hand_heart / hand_heart2) — chest-level, lifted to overhead "
     "by an arm-elevation augmentation"),
    ("COCO 2017", "person_keypoints_train2017 + images.cocodataset.org",
     "in-the-wild photos, full scenes, many sports/action shots",
     "annotations CC BY 4.0", "t_pose, raise_right_hand, raise_left_hand; + idle negatives"),
    ("Roboflow Universe", "5 ASL detection projects (ece496-public-asl, "
     "signlanguage-f0irs, actions-zqpb1, asl-detection-lvx6a, signlanguageai)",
     "frontal webcam ASL images, hand bbox labelled",
     "mostly CC BY 4.0", "i_love_you (ASL ILY handshape) — 606 rows; see §7"),
]

# (class label shown, samples_ext folder, caption, filename prefix | "")
SAMPLE_ORDER = [
    ("idle", "idle_hagrid", "HaGRID `no_gesture` — hand visible, no gesture", ""),
    ("idle", "idle_coco", "COCO — a person, arms down, no gesture", "c"),
    ("ok", "ok", "HaGRID `ok`", ""),
    ("two_finger", "two_finger", "HaGRID `peace` / `peace_inverted` / `two_up`", ""),
    ("thumb", "thumb", "HaGRID `like` — held-out HaGRID subjects are the thumb test set", ""),
    ("mini_heart", "mini_heart", "s01 `mini_heart_01` — HaGRIDv2 `hand_heart` trains it; "
                                 "an arm-elevation aug lifts those chest-level hands overhead", "s01"),
    ("raise_right_hand", "raise_right_hand", "COCO — person's right wrist above the nose", "c"),
    ("raise_left_hand", "raise_left_hand", "COCO + mirror-augmented raise_right_hand", "c"),
    ("t_pose", "t_pose", "COCO — both wrists ~shoulder height, arms spread wide", "c"),
    ("sit", "sit", "s01 `sit_01/02` test footage — TRAINED on 1,372 COCO rows "
                   "(torso upright, thighs level); COCO's are noisier (§9)", "s01"),
    ("laying", "laying", "s01 `laying_03` test footage — TRAINED on 607 COCO rows "
                         "(torso horizontal)", "s01"),
    ("squat", "squat", "s01 `squat_01` — COCO's squats were shallow crouches (§9), aug(s01) only", "s01"),
    ("glico_pose", "glico_pose", "s01 `glico_pose_01` — no public dataset, aug(s01) only", "s01"),
    ("(negative)", "ily_negative_call", "HaGRID `call` (shaka) -> mapped to idle "
                                        "(a distinctive non-target gesture, a good hard negative)", ""),
]


def _class_grid(folder: str, prefix: str, n: int = 4, tile: int = 210) -> Path | None:
    """2-row grid for one class: raw source images (top) over their MediaPipe
    overlays (bottom). Built from data/samples_ext/<folder>/."""
    d = SAMPLES / folder
    if not d.is_dir():
        return None
    raws = [p for p in sorted(d.glob(f"{prefix}*.jpg")) if "_pose" not in p.name]
    if not raws:                                    # fall back to any sample
        raws = [p for p in sorted(d.glob("*.jpg")) if "_pose" not in p.name]
    raws = raws[:n]
    if not raws:
        return None
    cols = []
    for raw in raws:
        ov = raw.with_name(raw.stem + "_pose.jpg")
        a = cv2.imread(str(raw))
        b = cv2.imread(str(ov)) if ov.exists() else a
        if a is None:
            continue
        a = cv2.resize(a, (tile, tile))
        b = cv2.resize(b, (tile, tile)) if b is not None else np.zeros_like(a)
        cols.append(np.vstack([a, b]))
    if not cols:
        return None
    grid = np.hstack(cols)
    GRIDS.mkdir(parents=True, exist_ok=True)
    out = GRIDS / f"{folder}.jpg"
    cv2.imwrite(str(out), grid)
    return out


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for c, h in enumerate(headers):
        run = t.cell(0, c).paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for r, row in enumerate(rows, 1):
        for c, v in enumerate(row):
            run = t.cell(r, c).paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)
    return t


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    card = json.loads((DS / "dataset_card.json").read_text())
    pc = card["per_class"]

    doc = Document()
    doc.add_heading("SKUBA gesture detection — Phase 2 dataset", 0)
    doc.add_paragraph(f"Generated {date.today().isoformat()}").runs[0].italic = True

    doc.add_paragraph(
        "We have only 2 recorded subjects (s01: all classes; s02: laying + sit). "
        "That is too few for a leak-free subject-wise split. Decision (RoboCup@Home "
        "context, licences not a blocker): TRAIN the classifier on public datasets "
        "and keep every original s01/s02 frame as a held-out TEST set. This report "
        "is what to tell the team: which datasets, what the images look like, and "
        "how the training set was built."
    )
    doc.add_paragraph(
        "Vocabulary: 12 classes. i_love_you and rock were cut in Phase 3 (MediaPipe "
        "Hands can't tell them apart from two_finger on our footage — §7); heart "
        "was cut (no dataset — §8). See CLAUDE.md."
    )

    doc.add_heading("1. What was downloaded", 1)
    _table(doc, ["Source", "Dataset", "Image style", "Licence", "Our classes"],
           [list(s) for s in SOURCES])
    doc.add_paragraph(
        "Nothing was downloaded in bulk. pipeline/extract_external.py streams the "
        "HuggingFace parquet shards one at a time (~50-500 MB, deleted right "
        "after), runs our MediaPipe Pose + Hands pipeline on each image, and keeps "
        "ONLY the 152-d normalised feature vector (~0.6 KB/frame). COCO downloads "
        "only the ~250 MB keypoint-annotation JSON plus the few hundred images "
        "that pass a pose filter. Roboflow projects are downloaded as COCO-format "
        "exports (needs a free API key), extracted, then deleted. No source "
        "images are retained. Extraction runs on Colab — the SKUBA laptop OOMs "
        "(docs/run_extraction_elsewhere.md)."
    )

    doc.add_heading("2. What the training images look like — every class", 1)
    doc.add_paragraph(
        "For each class: the source images (top row) over the MediaPipe skeleton "
        "+ wrist-anchored hand crop we actually turn into the 152-d feature "
        "vector (bottom row). This is exactly what the model sees — no pixels are "
        "kept, only the normalised keypoints."
    )
    for label, folder, cap, prefix in SAMPLE_ORDER:
        grid = _class_grid(folder, prefix)
        doc.add_paragraph(f"{label}   —   {cap}", style="Heading 3")
        if grid:
            doc.add_picture(str(grid), width=Inches(6.6))
        else:
            doc.add_paragraph(f"[no samples in data/samples_ext/{folder} — "
                              f"run `python scripts/pull_samples.py`]")
    doc.add_paragraph(
        "The COCO honesty problem is visible above: COCO has ACTIONS, not held "
        "gestures. 'arm raised' catches pitchers and riders; 't_pose' catches "
        "anyone mid-stride with arms out. The _classify_coco filter + MediaPipe "
        "re-verification (§4) drop ~40-55% of the COCO rows, but the survivors "
        "still skew sporty — this is why t_pose / raise_hand test ~0.75-0.84 and "
        "why Phase 6 field data matters. sit and laying mine cleanly; squat did "
        "not (§9)."
    )

    doc.add_heading("3. From image to feature", 1)
    for t in [
        "person -> MediaPipe Pose -> 33 body landmarks.",
        "each wrist -> square crop sized by shoulder width -> MediaPipe Hands -> "
        "21 landmarks/hand (or absent -> presence flag = 0).",
        "normalise: body centred on the shoulder-hip midpoint / scaled by "
        "shoulder-hip distance; each hand centred on its wrist / scaled by palm "
        "width (features/normalize.py).",
        "concatenate -> 152-d vector: body[0:66], left hand[66:108] + flag[108], "
        "right hand[109:151] + flag[151] (features/schema.py).",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("4. Cleaning the external features", 1)
    doc.add_paragraph(
        "pipeline/build_dataset.py::_clean_external drops: (a) rows with any "
        "|feature| > 12 — normalisation blows up when the hips are out of a webcam "
        "frame; (b) hand-shape classes (ok/rock/thumb/two_finger/mini_heart) with "
        "no hand detected; (c) COCO pose rows whose MediaPipe keypoints don't "
        "actually show the pose (COCO's own annotation frequently points at a "
        "different person in a multi-person photo)."
    )
    cs = card["external_clean"]
    _table(doc, ["source file", "raw rows", "kept", "%"],
           [[k.split("__")[-1], v["raw"], v["kept"],
             f"{100 * v['kept'] // max(1, v['raw'])}"] for k, v in cs.items()])

    doc.add_heading("5. Building train / test", 1)
    doc.add_paragraph(
        f"train.npz = {card['train']['rows']:,} rows: cleaned external features + "
        f"augmentation (mirror+relabel, limb-length jitter, rotation, keypoint "
        f"dropout, coord noise — features/augment.py). Per-class cap ~7,000; "
        f"scarce classes boosted to ~2,200."
    )
    doc.add_paragraph(
        f"test.npz = {card['test']['rows']:,} rows, NEVER augmented: 1,367 "
        f"original s01/s02 frames + 214 held-out HaGRID `thumb` (s01 has no "
        f"thumb). s01/s02 never appear in training except as augmented copies for "
        f"the 6 aug-only classes."
    )
    rows = []
    for c in card["classes"]:
        r = pc.get(c, {})
        rows.append([c, r.get("eval_type", "-"), r.get("train_source", "-"),
                     r.get("train_rows", 0), r.get("test_rows", 0)])
    _table(doc, ["class", "eval type", "train source", "train rows", "test rows"], rows)

    doc.add_heading("6. What each class's test number will mean", 1)
    for t in [
        "cross_domain (idle, ok, two_finger, rock, thumb, mini_heart, "
        "raise_L/R_hand, t_pose): external train, our test — a REAL "
        "cross-subject + cross-domain generalisation number.",
        "held_out_external (thumb): cross-subject within HaGRID, same webcam "
        "domain.",
        "aug_only_leak (sit, squat, laying, glico_pose): no external data. "
        "Train = augmented s01 frames. For `sit`, s02's frames are still a clean "
        "cross-person check (aug pool is s01-only); for `laying` there is no s01 "
        "so it is a full leak. NOT generalisation numbers — Phase 6 is the gate.",
        "pending_data (i_love_you, heart): 0 training rows, see §7-8.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("7. i_love_you / rock — every data path tried, then CUT", 1)
    doc.add_paragraph(
        "i_love_you is the ASL 'I Love You' handshape (thumb + index + pinky "
        "extended); rock is the horns/mano-cornuta (index + pinky). Both were cut "
        "from the vocabulary in Phase 3. The data history for i_love_you, in order:"
    )
    _table(doc, ["attempt", "data", "result"], [
        ["aug(s01) only", "the s01 i_love_you_01 clip, augmented",
         "recall 0.00; dragged rock to 0.00 (s01's rock and ILY were performed "
         "near-identically in one session)"],
        ["Round-1 search", "WLASL / MS-ASL video, ASL alphabet image sets",
         "no 'I love you' gloss / only A-Z — declared 'none' (this was a scoping "
         "miss: the sign is a static handshape, not covered by alphabet sets)"],
        ["Round-2 search", "Roboflow Universe ASL detection datasets",
         "the ILY handshape IS a labelled class in several — found 5 usable "
         "(ece496-public-asl, signlanguage-f0irs, actions-zqpb1, "
         "asl-detection-lvx6a, signlanguageai), mostly CC BY 4.0"],
        ["roboflow_ily extraction", "606 real ILY images through our pipeline "
         "(committed: data/features_ext/roboflow_ily__i_love_you.npz)",
         "recall still 0.00; rock fell 0.78 -> 0.48"],
        ["crop reframing", "two-stage tight crop so the hand fills the frame",
         "no change to the features — ILY thumbPinkyAng 0.26 vs rock 0.28"],
    ])
    doc.add_paragraph(
        "Root cause (measured on s01 hand features): MediaPipe Hands reports "
        "every finger curled for i_love_you / rock / two_finger alike on s01's "
        "conversational-distance footage — idxCurl ~2.7 rad for all three "
        "(a straight finger is ~0.4). The three classes are the SAME feature "
        "vector; no training data or preprocessing separates identical inputs "
        "(a two-stage tight crop was tested — no change). Decision: cut "
        "i_love_you and rock, keep two_finger (it still works at ~0.72 because "
        "HaGRID anchors 'loose fist + upright body' to that label). Revisit both "
        "if a better hand model is adopted (Phase 4) — the 606 ILY feature rows "
        "are kept for that. See docs/phase3_baseline.md, docs/external_datasets.md R2.9."
    )

    doc.add_heading("8. heart — no dataset exists, CUT", 1)
    doc.add_paragraph(
        "The overhead two-arm heart is in no public dataset. It is a pose-only "
        "class (arms make the shape), so it was treated like t_pose and a "
        "COCO-mining filter was added to _classify_coco (both wrists above the "
        "eyes, hands near the midline, both elbows outboard). Scanning all of "
        "COCO-train it finds only 55 candidates, most of them false positives "
        "(reaching / diving / celebration poses). MediaPipe would keep ~half — "
        "too few, and the i_love_you runs showed a class this small poisons its "
        "neighbours. Cut; revisit with Phase 6 field recordings."
    )

    doc.add_heading("9. Body postures — de-leaked from COCO", 1)
    doc.add_paragraph(
        "sit / squat / laying were originally AUG_ONLY (train = augmented s01/s02, "
        "test = the same frames — a leaked 0.96-1.00) only because Phase 2 spent "
        "its effort on the hand classes. Body posture is where MediaPipe Pose is "
        "reliable and COCO has plenty, so _classify_coco gained sit / squat / "
        "laying branches (spine horizontal -> laying; knee-bend + hip-vs-knee "
        "height -> sit / squat), re-verified against MediaPipe's normalized body."
    )
    _table(doc, ["class", "COCO rows (clean)", "result"], [
        ["sit", "1,372", "leaked 0.96 / real 0.53 -> honest 0.91 F1, 0.97 "
                "cross-person. Best single Phase 3 improvement."],
        ["laying", "607", "honest 0.71 (RF) / 0.88 (LGBM), was a leaked 1.00. "
                   "Misses: a lying arm reads as raise_left_hand. Phase 4/5."],
        ["squat", "919, but bad", "COCO 'squat' = shallow crouches (knee ~107 "
                  "deg), overlaps sit — with it in, squat F1 -> 0.01 and sit -> "
                  "0.58. Reverted to aug(s01)-only. Needs NTU / a gym set."],
    ])
    doc.add_paragraph(
        "The COCO `sit` / `laying` auto-labels are noisy — an athletic crouch or "
        "a bent-knee reach passes the 2D geometric filter, and COCO is mostly "
        "sports photos, not people sitting on chairs. ~9% of the geometrically-"
        "clean rows are still not really the posture. The aggregate signal "
        "survives (sit 0.91), but this is why sit is not 0.99 and why the sample "
        "grid in §2 shows the clean s01 footage rather than the training rows. "
        "glico_pose has no dataset — stays aug-only."
    )

    doc.add_heading("10. Known gaps", 1)
    for t in [
        "glico_pose — no external data; train = aug(s01).",
        "i_love_you / rock / heart — cut (§7-8).",
        "Hand-landmark resolution: MediaPipe Hands can't resolve fine finger "
        "differences on s01's distant footage — caps ok / two_finger at "
        "~0.72-0.78; a better hand model is the Phase 4 lever.",
        "COCO poses are action photos, not held gestures — t_pose / raise_hand "
        "test lower for this reason; squat auto-labels are noisy.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    out = OUT / "dataset_report.docx"
    doc.save(str(out))
    print(f"-> {out}")


if __name__ == "__main__":
    main()
