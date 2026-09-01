"""Build results/phase4/classifier_report.docx — the Track C model-family
comparison: N candidate classifiers on the finalised 12-class dataset, all
scored by the same harness.

    python -m classifier.train --model all
    python -m classifier.evaluate --all
    python scripts/phase4_report.py

Reads results/phase4/model_comparison.json + the eval jsons in
classifier/models/ + the figures in results/phase4/fig/ and results/phase3/fig/.
"""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
DS = ROOT / "data" / "dataset"
MODELS = ROOT / "classifier" / "models"
OUT = ROOT / "results" / "phase4"
FIG = OUT / "fig"
P3FIG = ROOT / "results" / "phase3" / "fig"

GREEN = RGBColor(0x1A, 0x7F, 0x37)
AMBER = RGBColor(0x88, 0x66, 0x00)

# nicer display order (best-known first)
ORDER = ["lgbm", "catboost", "hgb", "rf", "et", "svm", "mlp", "logreg"]
HEADLINE_ROWS = [
    ("macro_f1", "macro-F1 (12)"),
    ("macro_f1_real_eval", "macro-F1 (real-eval)"),
    ("micro_f1_accuracy", "accuracy"),
    ("balanced_accuracy", "balanced acc"),
    ("expected_calibration_error", "ECE (lower better)"),
    ("mean_conf_correct", "conf | correct"),
    ("mean_conf_wrong", "conf | wrong"),
    ("model_size_mb", "size (MB)"),
    ("fit_seconds", "fit (s)"),
]


def _cell(t, r, c, text, bold=False, color=None):
    cell = t.cell(r, c)
    cell.text = ""
    run = cell.paragraphs[0].add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color


def _table(doc, headers, rows, row_colors=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for c, h in enumerate(headers):
        _cell(t, 0, c, h, bold=True)
    for r, row in enumerate(rows, 1):
        for c, v in enumerate(row):
            _cell(t, r, c, v, color=(row_colors[r - 1] if row_colors else None))
    return t


def _fig(doc, path: Path, width=6.4):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
    else:
        doc.add_paragraph(f"[missing figure {path.name}]")


def main():
    cmp_path = OUT / "model_comparison.json"
    if not cmp_path.exists():
        raise SystemExit("run `python -m classifier.evaluate --all` first")
    cmp = json.loads(cmp_path.read_text())
    card = json.loads((DS / "dataset_card.json").read_text())
    classes = card["classes"]

    models = [m for m in ORDER if m in cmp["headline"]] + \
             [m for m in cmp["headline"] if m not in ORDER]
    hl = cmp["headline"]
    pcf1 = cmp["per_class_f1"]
    xper = cmp.get("cross_person", {})

    # winner = best macro_f1_real_eval, tie-broken by smaller size
    winner = min(models, key=lambda m: (-(hl[m].get("macro_f1_real_eval") or 0),
                                        hl[m].get("model_size_mb") or 1e9))

    doc = Document()
    doc.add_heading("SKUBA gesture detection — Phase 4: classifier comparison", 0)
    doc.add_paragraph(f"Generated {date.today().isoformat()}").runs[0].italic = True

    p = doc.add_paragraph()
    r = p.add_run(f"{len(models)} candidate classifiers trained on the same "
                  f"12-class pipeline (train {card['train']['rows']:,} / val "
                  f"{card.get('val', {}).get('rows', 0):,} / test "
                  f"{card['test']['rows']:,}), scored by classifier/evaluate.py. "
                  f"Leading model on the real-generalisation metric: "
                  f"{winner.upper()}.")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = GREEN

    doc.add_heading("1. Headline metrics — every candidate", 1)
    rows = []
    for key, label in HEADLINE_ROWS:
        rows.append([label] + [_fmt(hl[m].get(key), key) for m in models])
    _table(doc, ["metric"] + [m.upper() for m in models], rows)
    doc.add_paragraph(
        "'real-eval' = macro-F1 over the classes with a genuine cross-domain / "
        "cross-subject test (excludes the aug-only-leak classes squat, "
        "glico_pose). 'conf | correct' vs 'conf | wrong' is the mean max-"
        "probability when the prediction is right vs wrong — the gap is what the "
        "Phase 5 idle/unknown threshold needs; a model that is accurate but "
        "piles all its confidence at ~1.0 needs post-hoc calibration first."
    )
    _fig(doc, FIG / "compare_overall.png")

    doc.add_heading("2. Per-class F1", 1)
    prows, pcolors = [], []
    for c in classes:
        vals = pcf1.get(c, {})
        if not any(vals.get(m) is not None for m in models):
            continue
        best = max((vals.get(m) or 0) for m in models)
        prows.append([c] + [f"{vals.get(m):.2f}" if vals.get(m) is not None else "-"
                            for m in models])
        pcolors.append(AMBER if best < 0.75 else None)
    _table(doc, ["class"] + [m.upper() for m in models], prows, pcolors)
    _fig(doc, FIG / "compare_f1.png")

    doc.add_heading("3. Cross-person (sit — the only multi-subject class)", 1)
    xrows = []
    for m in models:
        w = xper.get(m, {}).get("sit", {})
        if w:
            xrows.append([m.upper(), w.get("s01", "-"), w.get("s02", "-")])
    if xrows:
        _table(doc, ["model", "sit @ s01", "sit @ s02 (never trained)"], xrows)
        doc.add_paragraph("s02 recall is the single most trustworthy number — a "
                          "person who never appears in training.")

    doc.add_heading("4. The leading model in detail — " + winner.upper(), 1)
    ev_path = MODELS / f"{winner}.eval.json"
    if ev_path.exists():
        ev = json.loads(ev_path.read_text())
        drows, dcolors = [], []
        for c, rr in ev["per_class"].items():
            drows.append([c, rr["eval_type"], rr["n"], rr["precision"],
                          rr["recall"], rr["f1"], rr["mean_conf"]])
            dcolors.append(AMBER if rr["f1"] < 0.75 else None)
        _table(doc, ["class", "eval type", "n", "prec", "rec", "F1", "conf"],
               drows, dcolors)
    for name in (f"{winner}_confusion.png", f"{winner}_calibration.png"):
        f = FIG / name
        if not f.exists():
            f = P3FIG / name
        _fig(doc, f)

    doc.add_heading("5. Verdict + what's next", 1)
    doc.add_paragraph(
        f"{winner.upper()} is the working pick. Remaining Phase 4 work: "
        "(a) if {winner}'s confidence is piled at ~1.0, apply Platt / isotonic / "
        "temperature calibration and re-score; (b) shrink any tree model that is "
        "the pick and over-sized; (c) Track D — tune AugParams strength against "
        "the val split; (d) Track E — per-class confidence thresholds on the val "
        "split for the idle/unknown fallback. Then lock and write "
        "docs/phase4_baseline.md.".replace("{winner}", winner.upper())
    )

    out = OUT / "classifier_report.docx"
    try:
        doc.save(str(out))
    except PermissionError:
        out = OUT / "classifier_report.NEW.docx"
        doc.save(str(out))
        print("!! report open in Word — wrote", out.name)
    print(f"-> {out}  (winner: {winner})")


def _fmt(v, key):
    if v is None:
        return "-"
    if key == "model_size_mb":
        return f"{v:.0f}" if v >= 1 else "<1"
    if key == "fit_seconds":
        return f"{v:.0f}"
    return f"{v:.3f}"


if __name__ == "__main__":
    main()
