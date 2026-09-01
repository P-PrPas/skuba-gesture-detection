"""Build results/phase3/classifier_report.docx — the Phase 3 baseline: two
candidate classifiers (LightGBM, RandomForest) trained on the same pipeline,
scored by the same harness, compared head-to-head.

    python -m classifier.train --model rf   --features both
    python -m classifier.train --model lgbm --features both
    python -m classifier.evaluate --all        # metrics json + figures
    python scripts/phase3_report.py

Figures come from classifier/evaluate.py (results/phase3/fig/). This script only
lays them out with the tables.
"""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
DS = ROOT / "data" / "dataset"
MODELS = ROOT / "classifier" / "models"
OUT = ROOT / "results" / "phase3"
FIG = OUT / "fig"

GREEN = RGBColor(0x1A, 0x7F, 0x37)
AMBER = RGBColor(0x88, 0x66, 0x00)


def _cell(t, r, c, text, bold=False, color=None):
    cell = t.cell(r, c)
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color


def _table(doc, headers, rows, row_colors=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for c, h in enumerate(headers):
        _cell(t, 0, c, h, bold=True)
    for r, row in enumerate(rows, 1):
        for c, v in enumerate(row):
            _cell(t, r, c, v, color=(row_colors[r - 1] if row_colors else None))
    return t


def _fig(doc, name, width=6.4):
    p = FIG / name
    if p.exists():
        doc.add_picture(str(p), width=Inches(width))
    else:
        doc.add_paragraph(f"[missing figure {name} — run `python -m classifier.evaluate --all`]")


def main():
    card = json.loads((DS / "dataset_card.json").read_text())
    classes = card["classes"]
    ev = {m: json.loads((MODELS / f"{m}.eval.json").read_text())
          for m in ("rf", "lgbm") if (MODELS / f"{m}.eval.json").exists()}
    if not ev:
        raise SystemExit("run `python -m classifier.evaluate --all` first")
    cmp = json.loads((OUT / "model_comparison.json").read_text()) \
        if (OUT / "model_comparison.json").exists() else None
    rf, lg = ev.get("rf"), ev.get("lgbm")
    primary = rf or next(iter(ev.values()))
    nmod = len(primary["per_class"])
    nreal = sum(1 for c, r in primary["per_class"].items()
                if r["eval_type"] in ("cross_domain", "held_out_external"))

    doc = Document()
    doc.add_heading("SKUBA gesture detection — Phase 3 baseline classifier", 0)
    doc.add_paragraph(f"Generated {date.today().isoformat()}").runs[0].italic = True

    p = doc.add_paragraph()
    r = p.add_run(f"Phase 3 trains two candidate classifiers on the fixed feature "
                  f"pipeline and compares them with one reusable eval harness. "
                  f"{nmod} of 15 target classes are modelled; 3 were cut "
                  f"(i_love_you, rock, heart).")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = GREEN
    if rf and lg:
        gap = round(rf["mean_conf_correct"] - rf["mean_conf_wrong"], 2)
        doc.add_paragraph(
            f"Headline: RandomForest macro-F1 {rf['macro_f1']} vs LightGBM "
            f"{lg['macro_f1']} (over the {nreal} non-leaked classes: "
            f"{rf['macro_f1_real_eval']} vs {lg['macro_f1_real_eval']}). Accuracy "
            f"is close. RandomForest is the working default because its "
            f"confidence is usable for the Phase 5 idle/unknown threshold — mean "
            f"confidence {rf['mean_conf_correct']} when correct vs "
            f"{rf['mean_conf_wrong']} when wrong (a {gap} gap), and its "
            f"confidence histogram is spread. LightGBM piles ~80% of predictions "
            f"into the top confidence bin ({lg['mean_conf_correct']} correct / "
            f"{lg['mean_conf_wrong']} wrong) so a threshold can't separate them. "
            f"LightGBM is 13x smaller ({lg['model_size_mb']} MB vs "
            f"{rf['model_size_mb']} MB) and wins on `laying` and `sit`@s02 — it "
            f"stays a Phase 4 candidate (with probability calibration). The final "
            f"model lock (incl. an MLP) is a Phase 4 decision."
        )

    # ---------------------------------------------------------------- 1
    doc.add_heading("1. Architecture — one classifier slot", 1)
    doc.add_paragraph(
        "The pipeline has exactly ONE classifier slot: fused 152-d feature vector "
        "in -> (class, confidence) out -> temporal smoothing -> final label or "
        "idle. Phase 3 does NOT stack or ensemble; it picks one model for that "
        "slot. LightGBM and RandomForest are the two Phase 3 candidates (chosen "
        "for zero-VRAM CPU inference and fast iteration); Phase 4 adds an MLP and "
        "locks one."
    )
    doc.add_paragraph(
        f"Features: `{primary['features']}` — the 152-d normalised keypoints PLUS "
        "37 derived features (joint angles + length ratios + inter-wrist distance, "
        "features/derived.py). The derived features are scale-invariant. "
        "class_weight = balanced x {aug-only 0.35, idle 0.6}. Features clipped to "
        "+-10 (normalisation blows up when the hips leave the frame)."
    )

    # ---------------------------------------------------------------- 2
    doc.add_heading("2. Train / test data", 1)
    doc.add_paragraph(
        f"Train: {card['train']['rows']:,} rows — cleaned external features "
        f"(HaGRID, COCO, Roboflow) + augmentation. Test: {card['test']['rows']:,} "
        f"rows, NEVER augmented — original s01/s02 frames + 214 held-out HaGRID "
        f"`thumb`. Full dataset construction, sources and sample images: "
        f"results/phase2/dataset_report.docx."
    )
    et_rows = []
    for c in classes:
        pc = card["per_class"].get(c, {})
        if not pc or c not in primary["per_class"]:
            continue
        et_rows.append([c, pc.get("eval_type", "-"), pc.get("train_rows", 0),
                        pc.get("test_rows", 0), pc.get("train_source", "-")])
    _table(doc, ["class", "eval type", "train rows", "test rows", "train source"], et_rows)
    doc.add_paragraph(
        "eval type: cross_domain = external train + our test (a real cross-subject "
        "AND cross-domain number); held_out_external = cross-subject within HaGRID; "
        "aug_only_leak = train is augmented s01/s02, test is the same frames — "
        "NOT a generalisation number (`squat`, `glico_pose` only).", style="List Bullet")

    # ---------------------------------------------------------------- 3
    doc.add_heading("3. The eval harness (reused for every future model)", 1)
    doc.add_paragraph(
        "`classifier/evaluate.py` scores any bundle {clf, classes, clip, features} "
        "whose `clf` has `.predict_proba` + `.classes_`. Per model it writes "
        "`<m>.eval.json` and three figures; with `--all` it also writes "
        "`model_comparison.json` and the side-by-side figures below. Train a new "
        "candidate (MLP, another tree lib) into the same bundle shape and it is "
        "scored identically — same metrics, same plots."
    )
    _table(doc, ["the harness reports", "why"], [
        ["per-class precision / recall / F1 + support", "the basic score, per class"],
        ["macro-F1 (all) / macro-F1 (real-eval) / macro-F1 (cross-domain)",
         "headline, with and without the leaked classes"],
        ["micro-F1 (accuracy), balanced accuracy", "overall, imbalance-aware"],
        ["confusion matrix (counts + row-normalised) + heatmap PNG", "which class becomes which"],
        ["cross-person recall (any class seen for >1 subject)", "the honest generalisation signal"],
        ["mean confidence when correct vs wrong; 10-bin reliability + ECE",
         "is the confidence usable for the Phase 5 idle threshold?"],
        ["fit time, model file size, feature count", "deployment cost"],
    ])

    # ---------------------------------------------------------------- 4
    doc.add_heading("4. LightGBM vs RandomForest — head to head", 1)
    if rf and lg:
        _table(doc, ["metric", "RandomForest", "LightGBM", "better"], [
            [f"macro-F1 ({nmod} modelled)", rf["macro_f1"], lg["macro_f1"],
             "RF" if rf["macro_f1"] >= lg["macro_f1"] else "LGBM"],
            [f"macro-F1 ({nreal} real-eval)", rf["macro_f1_real_eval"], lg["macro_f1_real_eval"],
             "RF" if rf["macro_f1_real_eval"] >= lg["macro_f1_real_eval"] else "LGBM"],
            ["accuracy", rf["micro_f1_accuracy"], lg["micro_f1_accuracy"],
             "RF" if rf["micro_f1_accuracy"] >= lg["micro_f1_accuracy"] else "LGBM"],
            ["balanced accuracy", rf["balanced_accuracy"], lg["balanced_accuracy"],
             "RF" if rf["balanced_accuracy"] >= lg["balanced_accuracy"] else "LGBM"],
            ["sit @ s02 (clean cross-person)",
             cmp["cross_person"]["rf"]["sit"]["s02"], cmp["cross_person"]["lgbm"]["sit"]["s02"],
             "LGBM"] if cmp else ["sit @ s02", "-", "-", "-"],
            ["conf when CORRECT / when WRONG",
             f"{rf['mean_conf_correct']} / {rf['mean_conf_wrong']}",
             f"{lg['mean_conf_correct']} / {lg['mean_conf_wrong']}",
             "RF (wider gap -> thresholdable)"],
            ["calibration error (ECE, lower better)",
             rf["expected_calibration_error"], lg["expected_calibration_error"], "LGBM"],
            ["model file size", f"{rf['model_size_mb']} MB", f"{lg['model_size_mb']} MB", "LGBM (13x)"],
            ["fit time", f"{rf['fit_seconds']} s", f"{lg['fit_seconds']} s", "RF"],
        ])
        doc.add_heading("4a. Per-class F1", 2)
        prows, pcolors = [], []
        for c in classes:
            a = cmp["per_class_f1"].get(c, {}).get("rf") if cmp else rf["per_class"].get(c, {}).get("f1")
            b = cmp["per_class_f1"].get(c, {}).get("lgbm") if cmp else lg["per_class"].get(c, {}).get("f1")
            if a is None and b is None:
                continue
            diff = "" if a is None or b is None else ("RF" if a - b > 0.05 else
                                                      "LGBM" if b - a > 0.05 else "=")
            prows.append([c, a, b, diff])
            pcolors.append(AMBER if diff in ("RF", "LGBM") else None)
        _table(doc, ["class", "RF F1", "LGBM F1", "gap>0.05"], prows, pcolors)
        _fig(doc, "compare_f1.png")
        doc.add_paragraph(
            "RF wins mini_heart (0.75 vs 0.57 — the arm-elevation synthetic rows "
            "suit bagging) and sit@s01 / idle; LightGBM wins laying (0.88 vs 0.71) "
            "and two_finger. Everything else is within noise.")
        _fig(doc, "compare_overall.png")

        doc.add_heading("4b. Confusion matrices", 2)
        _fig(doc, "rf_confusion.png")
        _fig(doc, "lgbm_confusion.png")
        doc.add_paragraph(
            "Both: the dominant leak is gestures -> idle (a frame with a weak "
            "detection falls to idle). RF also has laying -> raise_left_hand / sit "
            "(a lying arm reads as raised in un-rotated normalized coords); "
            "LightGBM handles laying better.")

        doc.add_heading("4c. Confidence calibration — why RF is the default", 2)
        _fig(doc, "rf_calibration.png")
        _fig(doc, "lgbm_calibration.png")
        doc.add_paragraph(
            "This is the deciding difference. The Phase 5 fallback is 'if "
            "confidence < T -> output idle/unknown', so the confidence has to "
            "separate right from wrong. RandomForest's confidence histogram is "
            "spread across 0.2-1.0 and its mean confidence is "
            f"{rf['mean_conf_correct']} when correct vs {rf['mean_conf_wrong']} "
            "when wrong — a threshold works. LightGBM piles ~1000 of 1254 test "
            "frames into the top confidence bin (mean "
            f"{lg['mean_conf_correct']} correct / {lg['mean_conf_wrong']} wrong); "
            "its aggregate ECE is low only because most of those are right, but a "
            "threshold can't fish the wrong ones out. LightGBM would need "
            "probability calibration (Platt / isotonic) before it is usable in "
            "the slot — a Phase 4 experiment.")

    doc.add_heading("5. Per-class detail (RandomForest — the default)", 1)
    rows, colors = [], []
    for c in classes:
        pc = primary["per_class"].get(c)
        if not pc:
            continue
        rows.append([c, pc["eval_type"], pc["n"], pc["precision"], pc["recall"],
                     pc["f1"], pc["mean_conf"]])
        colors.append(AMBER if pc["f1"] < 0.75 else None)
    _table(doc, ["class", "eval type", "n", "prec", "rec", "F1", "mean conf"], rows, colors)
    _fig(doc, "rf_perclass.png")
    xp = primary.get("cross_person", {})
    if xp:
        doc.add_paragraph(
            "Cross-person (a class recorded for >1 subject): "
            + "; ".join(f"{c} " + ", ".join(f"{s} {v}" for s, v in w.items())
                        for c, w in xp.items())
            + ". `sit` is the only one — its s02 recall (a person never in "
            "training) is the single most trustworthy number in Phase 3.")
    doc.add_paragraph(
        "Weak spots (F1 < 0.75, amber): laying 0.71 (a lying arm reads as "
        "raise_left_hand — tighten the COCO filter / add a torso-horizontal "
        "emphasis), idle 0.63 (precision 0.47 — over-triggers; several gestures "
        "fall to it at ~0.7 recall; Phase 5 temporal smoothing + per-class "
        "thresholds recover most). raise_right_hand 0.77 and ok 0.78 are "
        "MediaPipe-limited (COCO action photos / distant hand).")

    doc.add_heading("6. How the 12-class set was reached", 1)
    for h, t in [
        ("mini_heart fixed (0.00 -> 0.75), no new data",
         "HaGRIDv2 hand_heart is a chest-level finger-heart; s01's mini_heart is "
         "hands-together overhead. The per-hand normalisation makes the handshape "
         "slice position-invariant, so only the body pose carried the gap. "
         "features/augment.raise_arms shifts both forearms + hand points up by a "
         "shared offset; an inter-wrist-distance derived feature gives a direct "
         "hands-together signal. 45-frame test — run variance is wide."),
        ("i_love_you + rock CUT",
         "MediaPipe Hands reports every finger curled for i_love_you / rock / "
         "two_finger alike on s01's conversational-distance footage — the same "
         "feature vector. Tried: aug(s01), 606 real Roboflow ILY images, a "
         "two-stage tight crop. None separated them; training i_love_you dragged "
         "rock 0.78 -> 0.48. two_finger is kept (HaGRID anchors it) and cutting "
         "rock took it 0.72 -> 0.97. Revisit with a better hand model (Phase 4)."),
        ("heart CUT",
         "No dataset for the overhead two-arm heart; the COCO-mining filter finds "
         "only 55 candidates, mostly false positives. Phase 6 field data."),
        ("sit + laying de-leaked from COCO",
         "Both were AUG_ONLY (a leaked 0.96-1.00). _classify_coco gained "
         "sit/squat/laying branches, re-verified against MediaPipe's normalized "
         "body. sit: leaked 0.96 / real 0.53 -> honest 0.91, cross-person 0.97. "
         "laying: honest 0.71-0.88. squat's COCO labels were shallow crouches "
         "that collapsed into sit (squat F1 -> 0.01) so squat reverted to "
         "aug(s01)-only — it still needs a real squat dataset."),
    ]:
        doc.add_heading(h, 3)
        doc.add_paragraph(t)

    doc.add_heading("7. Phase 4 — model exploration on this fixed pipeline", 1)
    for i, t in enumerate([
        "Add an MLP as the third candidate; lock one of {RF, LGBM, MLP} with a "
        "documented rationale, using this harness.",
        "Calibrate LightGBM's probabilities (Platt / isotonic) and re-check "
        "whether its size advantage then makes it the pick.",
        "Shrink the RandomForest — 419 MB is a deployment liability; max_depth / "
        "min_samples_leaf tuning should cut it hard with little accuracy loss.",
        "Compare feature sets: raw vs derived vs both (the --features flag exists).",
        "Evaluate a hand-landmark model swap (RTMPose-Hand) — ok / two_finger sit "
        "at ~0.78 and the cut i_love_you / rock are MediaPipe-Hands-limited; "
        "re-add i_love_you / rock if it works.",
        "Re-extract COCO with a person-bbox crop to lift raise_hand / t_pose recall.",
        "Find a real squat source (NTU / a gym dataset) to de-leak squat.",
        "Per-class confidence thresholds on a validation slice, then wire the "
        "idle/unknown fallback (carried into Phase 5).",
    ], 1):
        doc.add_paragraph(f"{i}. {t}", style="List Bullet")

    out = OUT / "classifier_report.docx"
    try:
        doc.save(str(out))
    except PermissionError:
        out = OUT / "classifier_report.NEW.docx"
        doc.save(str(out))
        print("!! classifier_report.docx open in Word — wrote", out.name)
    print(f"-> {out}")


if __name__ == "__main__":
    main()
